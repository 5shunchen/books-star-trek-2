#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
小说导出脚本 - 将 markdown 格式的章节导出为 txt、docx 和分卷 md 格式
"""

import re
from pathlib import Path
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn

# 配置
CHAPTERS_DIR = Path("chapters")
OUTPUT_DIR = Path("output")
NOVEL_TITLE = "星途：从修仙到宇宙主宰"
AUTHOR = "5shunchen"

# 卷目结构
VOLUMES = [
    {"name": "第一卷：绝境觉醒", "chapters": range(1, 61), "file": "vol1_绝境觉醒.md"},
    {"name": "第二卷：星际风云", "chapters": range(61, 121), "file": "vol2_星际风云.md"},
    {"name": "第三卷：万族争霸", "chapters": range(121, 136), "file": "vol3_万族争霸.md"},
]

def get_volume_info(chapter_num):
    """根据章节号获取卷信息"""
    for vol in VOLUMES:
        if chapter_num in vol["chapters"]:
            return vol
    return None

def get_chapter_files():
    """获取所有章节文件，按章节号排序"""
    chapter_files = []
    pattern = re.compile(r'第 (\d+) 章')

    for f in CHAPTERS_DIR.glob("第*.md"):
        match = pattern.search(f.name)
        if match:
            chapter_num = int(match.group(1))
            chapter_files.append((chapter_num, f))

    chapter_files.sort(key=lambda x: x[0])
    return chapter_files

def read_chapter(filepath):
    """读取章节内容，保留 markdown 格式"""
    with open(filepath, 'r', encoding='utf-8') as f:
        return f.read()

def read_chapter_plain(filepath):
    """读取章节内容，移除 markdown 格式"""
    with open(filepath, 'r', encoding='utf-8') as f:
        content = f.read()

    lines = content.split('\n')
    cleaned_lines = []

    for line in lines:
        line = re.sub(r'^#+\s*', '', line)
        line = line.replace('**', '')
        line = line.replace('*', '')
        cleaned_lines.append(line)

    return '\n'.join(cleaned_lines)

def export_txt(chapter_files, output_path):
    """导出为 txt 格式"""
    print(f"正在导出 TXT 格式，共{len(chapter_files)}章...")

    with open(output_path, 'w', encoding='utf-8') as f:
        f.write("=" * 60 + "\n")
        f.write(f"{NOVEL_TITLE}\n")
        f.write("=" * 60 + "\n\n")
        f.write(f"作者：{AUTHOR}\n")
        f.write(f"总章节数：{len(chapter_files)}\n")
        f.write("\n" + "=" * 60 + "\n\n")

        # 卷目结构
        f.write("卷目结构\n")
        f.write("-" * 60 + "\n")
        for vol in VOLUMES:
            chapter_list = list(vol["chapters"])
            f.write(f"{vol['name']} (第{chapter_list[0]}-{chapter_list[-1]}章)\n")
        f.write("\n" + "=" * 60 + "\n\n")

        # 目录
        f.write("目 录\n")
        f.write("-" * 60 + "\n")
        current_volume = ""
        for chapter_num, filepath in chapter_files:
            vol = get_volume_info(chapter_num)
            if vol and vol["name"] != current_volume:
                f.write(f"\n【{vol['name']}】\n")
                current_volume = vol["name"]
            chapter_name = filepath.stem.replace("（重构）", "")
            f.write(f"第{chapter_num:03d}章 {chapter_name}\n")
        f.write("\n" + "=" * 60 + "\n\n")

        # 正文
        current_volume = ""
        for i, (chapter_num, filepath) in enumerate(chapter_files):
            vol = get_volume_info(chapter_num)
            if vol and vol["name"] != current_volume:
                f.write(f"\n{'=' * 60}\n")
                f.write(f"{vol['name']}\n")
                f.write(f"{'=' * 60}\n\n")
                current_volume = vol["name"]

            chapter_name = filepath.stem.replace("（重构）", "")
            content = read_chapter_plain(filepath)

            f.write(f"\n{'-' * 60}\n")
            f.write(f"第{chapter_num}章 {chapter_name}\n")
            f.write(f"{'-' * 60}\n\n")
            f.write(content)
            f.write("\n\n")

            if (i + 1) % 10 == 0:
                print(f"  已处理 {i + 1}/{len(chapter_files)} 章")

    print(f"TXT 导出完成：{output_path}")
    return output_path

def export_docx(chapter_files, output_path):
    """导出为 docx 格式"""
    print(f"正在导出 DOCX 格式，共{len(chapter_files)}章...")

    doc = Document()

    style = doc.styles['Normal']
    style.font.name = 'SimSun'
    style.font.size = Pt(12)
    style._element.rPr.rFonts.set(qn('w:eastAsia'), 'SimSun')

    # 封面
    doc.add_heading(NOVEL_TITLE, 0)
    doc.add_paragraph(f"作者：{AUTHOR}", style='Subtitle')
    doc.add_paragraph(f"总章节数：{len(chapter_files)}", style='Subtitle')
    doc.add_page_break()

    # 卷目结构
    doc.add_heading('卷目结构', level=1)
    for vol in VOLUMES:
        chapter_list = list(vol["chapters"])
        p = doc.add_paragraph()
        p.add_run(f"· {vol['name']}").bold = True
        doc.add_paragraph(f"  第{chapter_list[0]}章 - 第{chapter_list[-1]}章", style='List Bullet')
    doc.add_page_break()

    # 目录
    doc.add_heading('目录', level=1)
    current_volume = ""
    for chapter_num, filepath in chapter_files:
        vol = get_volume_info(chapter_num)
        if vol and vol["name"] != current_volume:
            p = doc.add_paragraph()
            p.add_run(f"\n【{vol['name']}】").bold = True
            current_volume = vol["name"]
        chapter_name = filepath.stem.replace("（重构）", "")
        doc.add_paragraph(f"第{chapter_num:03d}章 {chapter_name}", style='List Bullet')
    doc.add_page_break()

    # 正文
    current_volume = ""
    for i, (chapter_num, filepath) in enumerate(chapter_files):
        vol = get_volume_info(chapter_num)
        if vol and vol["name"] != current_volume:
            vol_heading = doc.add_heading(vol["name"], level=1)
            vol_heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
            doc.add_page_break()
            current_volume = vol["name"]

        chapter_name = filepath.stem.replace("（重构）", "")
        content = read_chapter_plain(filepath)

        heading = doc.add_heading(f"第{chapter_num}章 {chapter_name}", level=2)
        heading.alignment = WD_ALIGN_PARAGRAPH.CENTER

        sections = content.split('\n\n')
        for section in sections:
            section = section.strip()
            if section:
                if section.startswith('第') and '节' in section and len(section) < 50:
                    p = doc.add_paragraph(section)
                    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    p.runs[0].bold = True
                else:
                    doc.add_paragraph(section)

        if (i + 1) % 10 == 0:
            print(f"  已处理 {i + 1}/{len(chapter_files)} 章")

    doc.save(output_path)
    print(f"DOCX 导出完成：{output_path}")
    return output_path

def export_volumes_md(chapter_files):
    """按卷导出 md 文件"""
    print("正在分卷导出 MD 格式...")

    volumes_content = {vol["file"]: [] for vol in VOLUMES}

    for chapter_num, filepath in chapter_files:
        vol = get_volume_info(chapter_num)
        if vol:
            chapter_name = filepath.stem.replace("（重构）", "")
            content = read_chapter(filepath)
            volumes_content[vol["file"]].append((chapter_num, chapter_name, content))

    for vol in VOLUMES:
        output_path = OUTPUT_DIR / vol["file"]
        with open(output_path, 'w', encoding='utf-8') as f:
            # 卷标题
            f.write(f"# {vol['name']}\n\n")
            f.write(f"**章节范围**: 第{list(vol['chapters'])[0]}章 - 第{list(vol['chapters'])[-1]}章\n\n")
            f.write(f"**作者**: {AUTHOR}\n\n")
            f.write("---\n\n")

            # 目录
            f.write("## 目录\n\n")
            for chapter_num, chapter_name, _ in volumes_content[vol["file"]]:
                # 创建锚点链接
                anchor = chapter_name.replace(' ', '-').replace(':', '')
                f.write(f"- [第{chapter_num}章 {chapter_name}](#第{chapter_num}章-{chapter_name.replace(' ', '-')})\n")
            f.write("\n---\n\n")

            # 正文
            for chapter_num, chapter_name, content in volumes_content[vol["file"]]:
                f.write(f"\n{content}\n\n")

        print(f"  已导出：{output_path}")

    return OUTPUT_DIR

def generate_readme_chapter_list():
    """生成 README 用的章节列表"""
    chapter_files = get_chapter_files()

    result = {}
    for vol in VOLUMES:
        result[vol["name"]] = []

    for chapter_num, filepath in chapter_files:
        vol = get_volume_info(chapter_num)
        if vol:
            chapter_name = filepath.stem.replace("（重构）", "")
            result[vol["name"]].append((chapter_num, chapter_name))

    return result

def main():
    """主函数"""
    OUTPUT_DIR.mkdir(exist_ok=True)

    chapter_files = get_chapter_files()
    print(f"找到 {len(chapter_files)} 个章节文件")

    if not chapter_files:
        print("错误：未找到章节文件")
        return

    # 导出三种格式
    txt_path = OUTPUT_DIR / f"{NOVEL_TITLE}.txt"
    export_txt(chapter_files, txt_path)

    docx_path = OUTPUT_DIR / f"{NOVEL_TITLE}.docx"
    export_docx(chapter_files, docx_path)

    export_volumes_md(chapter_files)

    # 生成章节列表供 README 使用
    chapter_list = generate_readme_chapter_list()
    print("\n章节列表（可用于 README）:")
    for vol_name, chapters in chapter_list.items():
        print(f"\n{vol_name}:")
        for num, name in chapters:
            print(f"  第{num}章 {name}")

    print("\n" + "=" * 60)
    print("导出完成!")
    print(f"  TXT:    {txt_path}")
    print(f"  DOCX:   {docx_path}")
    print(f"  MD 分卷：{OUTPUT_DIR}/vol*.md")

if __name__ == "__main__":
    main()
